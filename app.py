import streamlit as st
import anthropic
import re
import json
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="AI Architect Assistant",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Custom CSS ──────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600&family=Syne:wght@400;700;800&display=swap');

html, body, [class*="css"] { font-family: 'Syne', sans-serif; }
.stApp { background: #0a0a0f; color: #e8e8f0; }
#MainMenu, footer, header { visibility: hidden; }

.hero {
    background: linear-gradient(135deg, #0d0d1a 0%, #111128 50%, #0a0a12 100%);
    border: 1px solid #2a2a4a; border-radius: 16px;
    padding: 44px 40px 32px; margin-bottom: 28px;
    position: relative; overflow: hidden;
}
.hero::before {
    content: ''; position: absolute; top: -60px; right: -60px;
    width: 300px; height: 300px;
    background: radial-gradient(circle, rgba(99,102,241,0.15) 0%, transparent 70%);
    pointer-events: none;
}
.hero-tag {
    font-family: 'IBM Plex Mono', monospace; font-size: 11px;
    letter-spacing: 3px; color: #6366f1; text-transform: uppercase; margin-bottom: 12px;
}
.hero h1 {
    font-size: 2.6rem; font-weight: 800;
    background: linear-gradient(90deg, #e8e8f0 0%, #a5b4fc 100%);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
    margin: 0 0 10px; line-height: 1.1;
}
.hero p { color: #8888aa; font-size: 0.95rem; max-width: 640px; line-height: 1.6; margin: 0; }

.intent-badge {
    display: inline-flex; align-items: center; gap: 8px;
    background: rgba(99,102,241,0.1); border: 1px solid rgba(99,102,241,0.3);
    border-radius: 8px; padding: 8px 16px; margin-bottom: 20px;
    font-family: 'IBM Plex Mono', monospace; font-size: 12px; color: #a5b4fc;
}
.intent-badge .dot { width: 7px; height: 7px; border-radius: 50%; background: #6366f1; }

.mode-chips { display: flex; flex-wrap: wrap; gap: 8px; margin-bottom: 20px; }
.mode-chip {
    font-family: 'IBM Plex Mono', monospace; font-size: 11px;
    letter-spacing: 1px; padding: 5px 14px; border-radius: 20px;
    border: 1px solid #2a2a4a; color: #6666aa; background: #0f0f1e;
}

.stTextArea label {
    font-family: 'IBM Plex Mono', monospace; font-size: 12px;
    letter-spacing: 2px; color: #6366f1 !important; text-transform: uppercase;
}
.stTextArea textarea {
    background: #0f0f1e !important; border: 1px solid #2a2a4a !important;
    border-radius: 10px !important; color: #e8e8f0 !important;
    font-family: 'IBM Plex Mono', monospace !important; font-size: 14px !important;
    line-height: 1.7 !important; padding: 16px !important; transition: border-color 0.2s;
}
.stTextArea textarea:focus {
    border-color: #6366f1 !important;
    box-shadow: 0 0 0 2px rgba(99,102,241,0.15) !important;
}
.stButton > button {
    background: linear-gradient(135deg, #6366f1, #818cf8) !important;
    color: white !important; border: none !important; border-radius: 10px !important;
    font-family: 'IBM Plex Mono', monospace !important; font-size: 13px !important;
    font-weight: 600 !important; letter-spacing: 1.5px !important;
    padding: 14px 36px !important; text-transform: uppercase !important;
    transition: all 0.2s !important; width: 100%;
}
.stButton > button:hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 8px 24px rgba(99,102,241,0.35) !important;
}

.section-card {
    background: #0f0f1e; border: 1px solid #1e1e38; border-radius: 14px;
    padding: 28px 32px; margin-bottom: 20px; transition: border-color 0.2s;
}
.section-card:hover { border-color: #3a3a6a; }
.section-header {
    display: flex; align-items: center; gap: 14px;
    margin-bottom: 18px; padding-bottom: 14px; border-bottom: 1px solid #1e1e38;
}
.section-num {
    font-family: 'IBM Plex Mono', monospace; font-size: 11px; color: #6366f1;
    background: rgba(99,102,241,0.12); border: 1px solid rgba(99,102,241,0.25);
    border-radius: 6px; padding: 3px 10px; letter-spacing: 2px;
}
.section-title { font-size: 1.05rem; font-weight: 700; color: #c8c8e8; letter-spacing: 0.3px; }
.section-body {
    font-size: 0.92rem; line-height: 1.85; color: #a8a8c8;
    white-space: pre-wrap; font-family: 'IBM Plex Mono', monospace;
}
.ascii-block {
    background: #080812; border: 1px solid #2a2a4a; border-left: 3px solid #6366f1;
    border-radius: 8px; padding: 16px 20px; margin-top: 10px;
    font-family: 'IBM Plex Mono', monospace; font-size: 12px;
    color: #a5b4fc; overflow-x: auto; white-space: pre;
}
.freeform-card {
    background: #0f0f1e; border: 1px solid #1e1e38; border-radius: 14px;
    padding: 28px 32px; margin-bottom: 20px;
}
.freeform-body {
    font-size: 0.93rem; line-height: 1.85; color: #a8a8c8;
    white-space: pre-wrap; font-family: 'IBM Plex Mono', monospace;
}
.status-bar {
    background: rgba(99,102,241,0.08); border: 1px solid rgba(99,102,241,0.2);
    border-radius: 8px; padding: 10px 18px;
    font-family: 'IBM Plex Mono', monospace; font-size: 12px;
    color: #818cf8; margin-bottom: 24px;
}
.custom-divider { border: none; border-top: 1px solid #1e1e38; margin: 28px 0; }

/* Download button override — distinct teal style */
.download-wrap .stDownloadButton > button {
    background: linear-gradient(135deg, #0d9488, #14b8a6) !important;
    font-size: 12px !important; padding: 10px 24px !important;
    letter-spacing: 1px !important; width: auto !important;
}
.download-wrap .stDownloadButton > button:hover {
    box-shadow: 0 6px 20px rgba(20,184,166,0.35) !important;
}
</style>
""", unsafe_allow_html=True)

# ── Constants ───────────────────────────────────────────────────────────────────
INTENTS = {
    "ARCHITECTURE_DESIGN": {"label": "Architecture Design", "icon": "🏗️", "desc": "6-section structured output"},
    "INTERVIEW_QA":         {"label": "Interview Q&A",       "icon": "🎯", "desc": "AI Architect interview prep"},
    "PROJECT_IDEAS":        {"label": "Project Ideas",       "icon": "💡", "desc": "AI Architect project suggestions"},
    "ARCHITECTURE_REVIEW":  {"label": "Architecture Review", "icon": "🔍", "desc": "Critique & improvement advice"},
    "TECH_STACK":           {"label": "Tech Stack",          "icon": "⚙️", "desc": "Tool & framework recommendations"},
    "UNKNOWN":              {"label": "General",             "icon": "💬", "desc": "General AI architecture guidance"},
}

ARCH_SECTIONS = [
    ("01", "Requirements & Constraints Summary", "REQUIREMENTS_AND_CONSTRAINTS"),
    ("02", "Architecture Choice + Justification", "ARCHITECTURE_CHOICE"),
    ("03", "Component Diagram",                   "COMPONENT_DIAGRAM"),
    ("04", "Top 3 Failure Modes + Mitigation",    "FAILURE_MODES"),
    ("05", "Tool Design",                          "TOOL_DESIGN"),
    ("06", "Tradeoffs Accepted",                   "TRADEOFFS"),
]

# ── System prompts ──────────────────────────────────────────────────────────────
ROUTER_PROMPT = """You are an intent classifier for an AI Architecture Assistant.
Classify the user's query into exactly one of:
ARCHITECTURE_DESIGN, INTERVIEW_QA, PROJECT_IDEAS, ARCHITECTURE_REVIEW, TECH_STACK, UNKNOWN
Respond ONLY with valid JSON: {"intent": "INTENT_NAME", "confidence": 0.95, "reason": "one sentence"}"""

ARCH_DESIGN_PROMPT = """You are a senior AI systems architect. Produce a rigorous architectural analysis using EXACTLY these markers:

===REQUIREMENTS_AND_CONSTRAINTS===
[Restate requirements. Name constraints explicitly. Include assumptions.]

===ARCHITECTURE_CHOICE===
[Name the architecture. Justify against constraints. Include 2-3 ruled-out alternatives.]

===COMPONENT_DIAGRAM===
[ASCII diagram: LLMs, retrieval, tools, validation, escalation, data flows. Label everything.]

===FAILURE_MODES===
[3 failure modes. For each: what breaks, why, specific mitigation. Be precise.]

===TOOL_DESIGN===
[3-5 core tools: Name, Description, Input schema with types, Output schema with types, Error handling.]

===TRADEOFFS===
[Name what you gave up. "Traded X for Y because Z." What you'd revisit if a constraint changed.]"""

INTERVIEW_PROMPT = """You are a senior AI Architect with 10+ years experience, conducting interview prep.
Give realistic questions with model answers, what interviewers look for, pitfalls, and level-specific tips (startup/mid-size/FAANG). Be specific and opinionated."""

PROJECT_IDEAS_PROMPT = """You are a senior AI Architect and mentor. Generate specific, actionable project ideas.
For each: name, one-line pitch, core technical challenge, key components, why it proves AI Architect skills, realistic scope, where to go deeper."""

REVIEW_PROMPT = """You are a principal AI systems architect doing a rigorous design review.
Structure: 1) What's solid, 2) Risks & gaps, 3) Failure modes, 4) Concrete improvements, 5) Questions to resolve.
Be honest and direct."""

TECHSTACK_PROMPT = """You are a senior AI systems architect. Give opinionated tech stack recommendations.
For each: what to use (specific, not categories), why, when NOT to use it, alternatives.
Cover: models, orchestration, retrieval, storage, observability, deployment, evaluation. Be current (2025 ecosystem)."""

UNKNOWN_PROMPT = """You are a senior AI systems architect and mentor. Answer with depth and specificity.
Be direct and opinionated. Provide concrete examples."""

INTENT_PROMPTS = {
    "ARCHITECTURE_DESIGN": ARCH_DESIGN_PROMPT,
    "INTERVIEW_QA": INTERVIEW_PROMPT,
    "PROJECT_IDEAS": PROJECT_IDEAS_PROMPT,
    "ARCHITECTURE_REVIEW": REVIEW_PROMPT,
    "TECH_STACK": TECHSTACK_PROMPT,
    "UNKNOWN": UNKNOWN_PROMPT,
}

EXAMPLES = [
    ("🏗️ Architecture", "Build a RAG-based customer support bot for a health insurance company. Must cite sources, handle 10k queries/day, never hallucinate coverage details, escalate edge cases to humans. Latency < 4s."),
    ("🎯 Interview",    "What are the top 5 system design interview questions for a senior AI Architect role? Give me model answers."),
    ("💡 Projects",     "I'm a backend engineer moving into AI. What are 3 portfolio projects that would prove I can design production AI systems?"),
    ("🔍 Review",       "I built a multi-agent system: agent A searches the web, agent B summarizes, agent C writes the report. They communicate via a shared Redis queue. What are the problems?"),
    ("⚙️ Tech Stack",   "What's the best stack for a production RAG system in 2025? I need retrieval, reranking, tracing, and eval. Team of 3 engineers."),
]

# ── LLM helpers ─────────────────────────────────────────────────────────────────
def classify_intent(query: str, client) -> dict:
    msg = client.messages.create(
        model="claude-sonnet-4-20250514", max_tokens=200,
        system=ROUTER_PROMPT,
        messages=[{"role": "user", "content": query}],
    )
    try:
        return json.loads(msg.content[0].text.strip())
    except Exception:
        return {"intent": "UNKNOWN", "confidence": 0.5, "reason": "Could not parse intent"}

def call_llm(query: str, system: str, client) -> str:
    msg = client.messages.create(
        model="claude-sonnet-4-20250514", max_tokens=4096,
        system=system,
        messages=[{"role": "user", "content": query}],
    )
    return msg.content[0].text

def parse_arch_sections(raw: str) -> dict:
    result = {}
    for _, _, marker in ARCH_SECTIONS:
        pattern = rf"==={marker}===\s*(.*?)(?====\w+===|$)"
        match = re.search(pattern, raw, re.DOTALL)
        result[marker] = match.group(1).strip() if match else ""
    return result

# ── DOCX builders ───────────────────────────────────────────────────────────────
ACCENT   = RGBColor(0x63, 0x66, 0xF1)   # indigo
DARK_BG  = RGBColor(0x1E, 0x1E, 0x38)
BODY_CLR = RGBColor(0x33, 0x33, 0x55)
TEXT_CLR = RGBColor(0x1A, 0x1A, 0x2E)

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background color via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _para_border_bottom(para, hex_color="6366F1", size=12):
    """Add a bottom border to a paragraph (used for section headings)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def _add_section_heading(doc: Document, num: str, title: str):
    """Styled section heading with accent underline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    _para_border_bottom(p)
    run_num = p.add_run(f"{num}  ")
    run_num.font.size  = Pt(10)
    run_num.font.color.rgb = ACCENT
    run_num.font.bold  = True
    run_num.font.name  = "Courier New"
    run_title = p.add_run(title.upper())
    run_title.font.size  = Pt(13)
    run_title.font.bold  = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x3E)
    run_title.font.name  = "Arial"

def _add_body_text(doc: Document, text: str, monospace: bool = False):
    """Add multi-line body text, preserving blank lines."""
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.name = "Courier New" if monospace else "Arial"
        if not monospace:
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x44)

def _add_code_block(doc: Document, text: str):
    """Shaded monospace block for ASCII diagrams."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, "F0F0FA")
    cell.width = Inches(6.5)
    # Clear default paragraph, add lines
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for line in text.split("\n"):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(line)
        run.font.name  = "Courier New"
        run.font.size  = Pt(8.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x88)
    doc.add_paragraph()  # spacer after table

def _add_cover(doc: Document, query: str, intent_label: str, intent_icon: str):
    """Cover page."""
    # Title block
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tag = p_tag.add_run("AI ARCHITECT ASSISTANT")
    run_tag.font.name  = "Arial"
    run_tag.font.size  = Pt(9)
    run_tag.font.color.rgb = ACCENT
    run_tag.font.bold  = True

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(8)
    run_title = p_title.add_run("Analysis Report")
    run_title.font.name  = "Arial"
    run_title.font.size  = Pt(26)
    run_title.font.bold  = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x3E)

    # Intent badge
    p_badge = doc.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_badge.paragraph_format.space_before = Pt(6)
    run_badge = p_badge.add_run(f"{intent_icon}  {intent_label}")
    run_badge.font.name  = "Arial"
    run_badge.font.size  = Pt(11)
    run_badge.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(4)
    run_date = p_date.add_run(datetime.now().strftime("%B %d, %Y"))
    run_date.font.name  = "Arial"
    run_date.font.size  = Pt(10)
    run_date.font.color.rgb = RGBColor(0x88, 0x88, 0xAA)

    # Query box
    doc.add_paragraph().paragraph_format.space_before = Pt(16)
    p_ql = doc.add_paragraph()
    r = p_ql.add_run("QUERY")
    r.font.name = "Courier New"; r.font.size = Pt(8); r.font.color.rgb = ACCENT; r.font.bold = True

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, "EEEEF8")
    p = cell.paragraphs[0]
    run = p.add_run(query)
    run.font.name = "Arial"; run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x44)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

    doc.add_page_break()

def build_arch_docx(query: str, raw: str, intent_label: str, intent_icon: str) -> bytes:
    """Build a Word doc for Architecture Design (6-section structured output)."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = section.right_margin = Inches(1)
        section.top_margin  = section.bottom_margin = Inches(1)

    _add_cover(doc, query, intent_label, intent_icon)

    sections = parse_arch_sections(raw)
    for num, title, marker in ARCH_SECTIONS:
        content = sections.get(marker, "")
        if not content:
            continue
        _add_section_heading(doc, num, title)
        if marker == "COMPONENT_DIAGRAM":
            _add_code_block(doc, content)
        else:
            _add_body_text(doc, content)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def build_freeform_docx(query: str, raw: str, intent_label: str, intent_icon: str) -> bytes:
    """Build a Word doc for free-form responses."""
    doc = Document()

    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = section.right_margin = Inches(1)
        section.top_margin  = section.bottom_margin = Inches(1)

    _add_cover(doc, query, intent_label, intent_icon)

    # Single heading + body
    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_before = Pt(8)
    _para_border_bottom(p_h)
    r = p_h.add_run(f"{intent_icon}  {intent_label.upper()}")
    r.font.name = "Arial"; r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x3E)

    _add_body_text(doc, raw)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ── Render helpers ───────────────────────────────────────────────────────────────
def render_arch_output(raw: str):
    sections = parse_arch_sections(raw)
    st.markdown('<div class="status-bar">✓ Architecture analysis complete — 6 sections generated</div>', unsafe_allow_html=True)
    for num, title, marker in ARCH_SECTIONS:
        content = sections.get(marker, "")
        if not content:
            continue
        body_html = (f'<div class="ascii-block">{content}</div>'
                     if marker == "COMPONENT_DIAGRAM"
                     else f'<div class="section-body">{content}</div>')
        st.markdown(f"""
        <div class="section-card">
            <div class="section-header">
                <span class="section-num">{num}</span>
                <span class="section-title">{title}</span>
            </div>
            {body_html}
        </div>
        """, unsafe_allow_html=True)

def render_freeform_output(raw: str, intent_key: str):
    cfg = INTENTS[intent_key]
    st.markdown(f'<div class="status-bar">✓ {cfg["icon"]} {cfg["label"]} response generated</div>', unsafe_allow_html=True)
    st.markdown(f"""
    <div class="freeform-card">
        <div class="freeform-body">{raw}</div>
    </div>
    """, unsafe_allow_html=True)

def render_download_button(docx_bytes: bytes, filename: str):
    st.markdown('<div class="download-wrap">', unsafe_allow_html=True)
    st.download_button(
        label="⬇  DOWNLOAD AS WORD DOCUMENT",
        data=docx_bytes,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=False,
    )
    st.markdown('</div>', unsafe_allow_html=True)

# ── Sidebar ─────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ Configuration")
    api_key = st.text_input("Anthropic API Key", type="password", placeholder="sk-ant-...")
    st.markdown("---")
    st.markdown("**Model:** claude-sonnet-4-20250514")
    st.markdown("**Router:** auto-classifies intent")
    st.markdown("---")
    st.markdown("**Supported modes:**")
    for k, v in INTENTS.items():
        if k != "UNKNOWN":
            st.markdown(f"{v['icon']} {v['label']}")
    st.markdown("---")
    st.markdown("*Just type naturally — intent is auto-detected.*")

# ── Hero ────────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="hero">
    <div class="hero-tag">⬡ AI Systems Design Assistant</div>
    <h1>AI Architect Assistant</h1>
    <p>Design systems, prep for interviews, get project ideas, review existing designs, or get tech stack recommendations. Intent is detected automatically — and every response can be downloaded as a Word document.</p>
</div>
""", unsafe_allow_html=True)

chips_html = '<div class="mode-chips">' + "".join(
    f'<span class="mode-chip">{v["icon"]} {v["label"]}</span>'
    for k, v in INTENTS.items() if k != "UNKNOWN"
) + '</div>'
st.markdown(chips_html, unsafe_allow_html=True)

# ── Examples ────────────────────────────────────────────────────────────────────
st.markdown("**Try an example:**")
cols = st.columns(len(EXAMPLES))
selected_example = ""
for i, (col, (label, ex)) in enumerate(zip(cols, EXAMPLES)):
    with col:
        if st.button(label, key=f"ex_{i}"):
            selected_example = ex

# ── Input ───────────────────────────────────────────────────────────────────────
query = st.text_area(
    "YOUR QUERY",
    value=selected_example,
    height=150,
    placeholder="e.g. 'Design a multi-agent research system...' or 'What interview questions should I expect for a senior AI Architect role?'",
)

st.markdown("")
run_btn = st.button("⬡  ANALYZE & RESPOND", use_container_width=True)

# ── Run ─────────────────────────────────────────────────────────────────────────
if run_btn:
    if not api_key:
        st.error("⚠️  Please enter your Anthropic API key in the sidebar.")
    elif not query.strip():
        st.error("⚠️  Please enter a query.")
    else:
        st.markdown('<hr class="custom-divider">', unsafe_allow_html=True)
        try:
            client = anthropic.Anthropic(api_key=api_key)

            with st.spinner("Detecting intent..."):
                intent_result = classify_intent(query.strip(), client)
                intent_key = intent_result.get("intent", "UNKNOWN")
                if intent_key not in INTENTS:
                    intent_key = "UNKNOWN"
                cfg = INTENTS[intent_key]

            st.markdown(f"""
            <div class="intent-badge">
                <span class="dot"></span>
                Detected: {cfg['icon']} <strong>{cfg['label']}</strong>
                &nbsp;·&nbsp; {intent_result.get('reason', '')}
            </div>
            """, unsafe_allow_html=True)

            with st.spinner("Generating response..."):
                raw = call_llm(query.strip(), INTENT_PROMPTS[intent_key], client)

            # Render
            if intent_key == "ARCHITECTURE_DESIGN":
                render_arch_output(raw)
                docx_bytes = build_arch_docx(query.strip(), raw, cfg["label"], cfg["icon"])
            else:
                render_freeform_output(raw, intent_key)
                docx_bytes = build_freeform_docx(query.strip(), raw, cfg["label"], cfg["icon"])

            # Download button
            st.markdown('<hr class="custom-divider">', unsafe_allow_html=True)
            slug = re.sub(r"[^a-z0-9]+", "_", intent_key.lower())
            ts   = datetime.now().strftime("%Y%m%d_%H%M")
            render_download_button(docx_bytes, f"ai_architect_{slug}_{ts}.docx")

            with st.expander("📄 Raw LLM output"):
                st.text(raw)

        except anthropic.AuthenticationError:
            st.error("❌ Invalid API key. Check your key at console.anthropic.com")
        except anthropic.RateLimitError:
            st.error("❌ Rate limit hit. Wait a moment and retry.")
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
